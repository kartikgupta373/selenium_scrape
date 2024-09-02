from selenium import webdriver
from selenium.webdriver.common.by import By
import time
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt

# Function to fetch links from a given URL
def get_links(url, base_url, seen_links):
    driver = webdriver.Chrome()
    driver.get(url)
    time.sleep(10)
    links = driver.find_elements(By.TAG_NAME, 'a')
    new_links = []

    for link in links:
        href = link.get_attribute('href')
        if href:
            if href.startswith(base_url) and href not in seen_links:
                seen_links.add(href)
                new_links.append(href)

    driver.quit()
    return new_links

# Function to fetch HTML content from a list of links          
def fetch_html_from_links(links):
    driver = webdriver.Chrome()
    html_data = []

    try:
        for link in links:
            driver.get(link)
            # time.sleep(5)  # Adjust the wait time as needed
            html_body = driver.find_element(By.TAG_NAME, 'body').get_attribute('innerHTML')
            html_data.append((link, html_body))
            print(f"Fetched HTML from {link}")
    except Exception as e:
        print(f"Error occurred: {str(e)}")
    finally:
        driver.quit()

    return html_data

# Function to extract data from HTML content
def extract_data(html):
    soup = BeautifulSoup(html, 'html.parser')
    data = {
        'title': soup.title.string if soup.title else 'No Title',
        'headings': [h.get_text() for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])],
        'paragraphs': [p.get_text() for p in soup.find_all('p')],
        'links': [a['href'] for a in soup.find_all('a', href=True)]
    }
    return data

# Function to write organized data to DOCX
def write_to_docx(data, filename):
    doc = Document()
    for item in data:
        # add paragraphs to the DOCX file
        for paragraph in item['paragraphs']:
            doc.add_paragraph(paragraph, style='BodyText')
        # add links to the DOCX file
        for link in item['links']:
            doc.add_paragraph(link, style='BodyText')

    doc.save(filename)

# Main function
def main():
    base_url = "https://www.orangemantra.com"
    seen_links = set()
    
    # Read existing links from the text file
    try:
        with open("fetched_links.txt", "r") as file:
            existing_links = set(line.strip() for line in file)
            seen_links.update(existing_links)
    except FileNotFoundError:
        existing_links = set()

    initial_links = get_links(base_url, base_url, seen_links)
    all_data = []

    # Fetch HTML and extract data from initial links
    html_data = fetch_html_from_links(initial_links)
    for link, html in html_data:
        data = extract_data(html)
        all_data.append(data)

        # Add new links found in the HTML content
        page_links = [link for link in data['links'] if link.startswith(base_url)]
        new_links = [link for link in page_links if link not in seen_links]
        seen_links.update(new_links)

    # Write new unique links to the text file
    with open("fetched_links.txt", "w") as file:
        for link in seen_links:
            file.write(link + "\n")

    # Write the organized data to a DOCX file   
    write_to_docx(all_data, "orangemantra_data.docx")
    print("Data written to doc file successfully!")

if __name__ == "__main__":
    main()
